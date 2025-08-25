import json
from langchain.docstore.document import Document
from docx import Document as DocxDocument
from langchain.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.document_loaders import PyPDFLoader

import os

current_dir = os.path.dirname(__file__) 
INPUT_FOLDER = os.path.join(current_dir, "data")
OUTPUT_VDB_PATH = os.path.join(current_dir, "faiss_index")
json_file_path = os.path.join(current_dir, "data/farming_expert.jsonl")

def load_jsonl_as_qa(path):
    qa_docs = []
    with open(path, "r", encoding="utf-8") as f:
        for line in f:
            try:
                record = json.loads(line.strip())
                user_msg = None
                assistant_msg = None
                messages = record.get("messages", [])
                if messages:
                    for msg in messages:
                        if isinstance(msg, dict):
                            if msg.get("role") == "user":
                                user_msg = msg.get("content")
                            elif msg.get("role") == "assistant":
                                assistant_msg = msg.get("content")

                if user_msg and assistant_msg:
                    text = f"Q: {user_msg}\nA: {assistant_msg}"
                    qa_docs.append(Document(page_content=text))
            except json.JSONDecodeError:
                # Skip lines that are not valid JSON
                continue
            except Exception as e:
                # Catch any other unexpected errors during processing
                print(f"An error occurred while processing a line: {e}")
                continue
    return qa_docs

def read_word_docs(path):
    docs = []
    docx_obj = DocxDocument(path)  
    for i, para in enumerate(docx_obj.paragraphs):
        text = para.text.strip()
        if text:  # skip empty paragraphs
            docs.append(Document(page_content=text, metadata={"source": "word", "paragraph": i}))

    for t_idx, table in enumerate(docx_obj.tables):
        table_text = []
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                table_text.append(" | ".join(row_text))
        if table_text:
            docs.append(Document(
                page_content="\n".join(table_text),
                metadata={"source": "word", "type": "table", "id": t_idx}
            ))

    return docs

def read_word_folder(folder_path):
    all_docs = []
    for file in os.listdir(folder_path):
        if file.endswith(".docx"):
            file_path = os.path.join(folder_path, file)
            word_docs = read_word_docs(file_path)
            all_docs.extend(word_docs)
    return all_docs

def read_pdf_docs(path):
    docs = []
    loader = PyPDFLoader(path)
    pdf_docs = loader.load()
    for i, d in enumerate(pdf_docs):
        docs.append(Document(
            page_content=d.page_content,
            metadata={"source": os.path.basename(path), "type": "pdf", "page": i}
        ))
    return docs


def read_pdf_folder(folder_path):
    all_docs = []
    for file in os.listdir(folder_path):
        if file.endswith(".pdf"):
            file_path = os.path.join(folder_path, file)
            all_docs.extend(read_pdf_docs(file_path))
    return all_docs


qa_docs = load_jsonl_as_qa(json_file_path)
word_docs = read_word_folder(INPUT_FOLDER)
pdf_docs = read_pdf_folder(INPUT_FOLDER)

all_docs = qa_docs + word_docs + pdf_docs

embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
vectorstore = FAISS.from_documents(all_docs, embeddings)
vectorstore.save_local(OUTPUT_VDB_PATH)
